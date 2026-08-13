from __future__ import annotations

import json
from pathlib import Path

from ..i18n import t


SUPPORTED_FORMATS = ("md", "docx", "pdf", "anki", "json")


def export_report(
    markdown_text: str,
    *,
    output_path: Path,
    fmt: str = "md",
    questions: list | None = None,
    locale: str = "zh-CN",
) -> tuple[bool, str]:
    """Export a report to the requested format. Returns (ok, message).

    Export failures NEVER affect the main learning flow: the Markdown is always
    available and the caller decides what to do with the failure message.
    """
    fmt = fmt.lower()
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"unsupported format {fmt!r}; expected {SUPPORTED_FORMATS}")
    output_path = Path(output_path)
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        if fmt == "md":
            output_path.write_text(markdown_text, encoding="utf-8")
        elif fmt == "json":
            output_path.write_text(
                json.dumps({"markdown": markdown_text}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        elif fmt == "docx":
            _export_docx(markdown_text, output_path)
        elif fmt == "pdf":
            _export_pdf(markdown_text, output_path)
        elif fmt == "anki":
            _export_anki(questions or [], output_path)
        return True, t(locale, "export.ok", path=str(output_path))
    except Exception as exc:  # noqa: BLE001 - export must not break the flow
        return False, t(locale, "export.failed", error=str(exc))


def _export_docx(markdown_text: str, output_path: Path) -> None:
    """DOCX export. Uses python-docx when available; otherwise raises so the
    caller reports a clean failure (main flow unaffected)."""
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx not installed (pip install -e .[docx])") from exc
    document = docx.Document()
    for line in markdown_text.splitlines():
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            document.add_paragraph(line.strip())
    document.save(str(output_path))


def _export_pdf(markdown_text: str, output_path: Path) -> None:
    """PDF export via reportlab when available; otherwise raises cleanly."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise RuntimeError("reportlab not installed for PDF export") from exc
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 50
    for line in markdown_text.splitlines():
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line[:100])
        y -= 14
    c.save()


def _export_anki(questions: list, output_path: Path) -> None:
    """Anki CSV export: Front,Back,Tags."""
    rows = ["Front,Back,Tags"]
    for q in questions:
        front = getattr(q, "question_text", q.get("question_text", "") if isinstance(q, dict) else "")
        back = getattr(q, "correct_answer", q.get("correct_answer", "") if isinstance(q, dict) else "")
        rows.append(f'"{str(front).replace(chr(34), chr(34)*2)}","{str(back).replace(chr(34), chr(34)*2)}","recallforge-skill"')
    output_path.write_text("\n".join(rows), encoding="utf-8")
