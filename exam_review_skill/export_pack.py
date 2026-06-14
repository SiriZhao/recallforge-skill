from __future__ import annotations

from pathlib import Path
from typing import Any


def write_markdown(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def export_docx(markdown_files: list[Path], output_path: Path, warnings: list[str]) -> None:
    try:
        import docx

        doc = docx.Document()
        for md in markdown_files:
            doc.add_heading(md.stem, level=1)
            for line in md.read_text(encoding="utf-8", errors="ignore").splitlines():
                if line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:].strip(), style="List Bullet")
                elif line.strip():
                    doc.add_paragraph(line.strip())
        doc.save(str(output_path))
    except Exception as exc:
        warnings.append(f"DOCX export unavailable: {exc}. 已保留 Markdown 输出。")


def export_pdf(markdown_files: list[Path], output_path: Path, warnings: list[str]) -> None:
    warnings.append("PDF export is optional and no PDF backend is configured; skipped gracefully.")


def export_anki_csv(questions: list[dict[str, Any]], output_path: Path) -> None:
    rows = ["Front,Back,Tags"]
    for q in questions:
        front = q.get("question_text", "").replace('"', '""')
        back = q.get("answer", "").replace('"', '""')
        rows.append(f'"{front}","{back}","exam-review-skill"')
    output_path.write_text("\n".join(rows), encoding="utf-8")
